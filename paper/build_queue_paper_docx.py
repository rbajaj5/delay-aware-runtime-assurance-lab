"""Build a rendered DOCX from the queue-aware coordination manuscript."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(__file__).with_name("queue-aware-coordination-draft.md")
OUTPUT = Path(__file__).with_name("queue-aware-coordination-draft.docx")
FIGURE = ROOT / "artifacts" / "20260802" / "helicopter_3d" / "helicopter_3d_threshold_sweep.png"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 99, 107)
TABLE_FILL = "F4F6F9"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color, bold=False):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_picture_alt_text(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, RGBColor(0, 0, 0))
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = doc.styles["Title"]
    set_style_font(title, "Calibri", 24, INK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.line_spacing = 1.0

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 11, MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = doc.styles["List Number"]
    set_style_font(list_style, "Calibri", 11, RGBColor(0, 0, 0))
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.194)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("QUEUE-AWARE COORDINATION | EXPLORATORY COURSE PROJECT")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(footer)


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=11)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=11, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=11, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=11)


def add_body_paragraph(doc, text, style=None, alignment=None):
    paragraph = doc.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    add_inline(paragraph, text)
    return paragraph


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            set_run_font(run, size=8.7, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table


def add_equation(doc, equation_lines):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(equation_lines):
        run = paragraph.add_run(line.strip())
        set_run_font(run, name="Cambria Math", size=10.5, italic=True)
        if index < len(equation_lines) - 1:
            run.add_break()
    return paragraph


def build():
    doc = Document()
    setup_document(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    current_heading = ""
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, line[2:].strip())
            index += 1
            continue
        if line.startswith("**Draft manuscript") or line.startswith("**Repository"):
            text = line.replace("**", "")
            paragraph = doc.add_paragraph(style="Subtitle")
            add_inline(paragraph, text)
            index += 1
            continue
        if line.startswith("**Artifact date"):
            text = line.replace("**", "")
            paragraph = doc.add_paragraph(style="Subtitle")
            add_inline(paragraph, text)
            paragraph.paragraph_format.space_after = Pt(12)
            index += 1
            continue
        if line.startswith("## "):
            current_heading = line[3:].strip()
            doc.add_heading(current_heading, level=1)
            index += 1
            continue
        if line.startswith("### "):
            current_heading = line[4:].strip()
            doc.add_heading(current_heading, level=2)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = parse_table(table_lines)
            widths = (
                [2400, 900, 1000, 750, 1050, 1600, 1660]
                if len(rows[0]) == 7
                else [1100, 1000, 1200, 1500, 1900, 2660]
            )
            add_table(doc, rows, widths)
            if current_heading == "4.2 Freshness threshold sweep" and FIGURE.exists():
                inline_shape = doc.add_picture(str(FIGURE), width=Inches(6.1))
                set_picture_alt_text(
                    inline_shape,
                    "Freshness threshold sweep",
                    "Line chart showing accepted and landed helicopters increasing with message-age threshold while conflict pair-steps rise sharply after threshold three.",
                )
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_before = Pt(2)
                caption.paragraph_format.space_after = Pt(10)
                run = caption.add_run("Figure 1. Freshness threshold sweep from the retained benchmark schedule.")
                set_run_font(run, size=9, color=MUTED, italic=True)
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, size=11, color=INK, bold=True, italic=True)
            index += 1
            continue
        if line.startswith("\\["):
            equation = []
            index += 1
            while index < len(lines) and not lines[index].startswith("\\]"):
                equation.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            add_equation(doc, equation)
            continue
        if re.match(r"^\d+\. ", line):
            while index < len(lines) and re.match(r"^\d+\. ", lines[index]):
                paragraph = doc.add_paragraph(style="List Number")
                add_inline(paragraph, re.sub(r"^\d+\. ", "", lines[index]))
                index += 1
            continue
        if line.startswith("##") or line.startswith("###"):
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines) and lines[index].strip() and not any(
            lines[index].startswith(prefix)
            for prefix in ("# ", "## ", "### ", "| ", "> ", "\\[", "1. ", "2. ", "3. ")
        ):
            paragraph_lines.append(lines[index])
            index += 1
        add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_lines))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
